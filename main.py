# imports
import os
import imp
import io
import datetime
import docx
import time

# define our clear function 
def clear(): 
  
    # for windows 
    if os.name == 'nt': 
        _ = os.system('cls') 
  
    # for mac and linux(here, os.name is 'posix') 
    else: 
        _ = os.system('clear')

# Logger
def Log(string):
    # Open Log File
    with open("AppLog.log", "a") as logFile:
        logFile.write('[' + str(datetime.datetime.now()) + '] ' +string + '\n')

# Main execution
class Main:
    # Var
    notValid = True

    while notValid:
        # Get valid module
        modName = input("Enter module (ext) filename (without .py): ")
        try:
            module = imp.load_source( modName, os.path.join('ext', modName + ".py") )
            notValid = False
        except:
            notValid = True

    modInit = getattr( module, "extInfo" )
    modNovelInfo = getattr( module, "novelDetails" )
    modChapterLink = getattr( module, "chapterLinks" )
    modChInfo = getattr( module, "getChapterInfo" )

    # Print module (extension) details
    extInfo = modInit("Package Info:\n")
    print( extInfo )
    Log( extInfo )

    # Get Webnovel File Name, and create it
    wnName = input("Get output File Name: ")
    wnFile = wnName + ".txt"
    outputFile = io.open(wnFile, "a", encoding="utf-8")
    # Create Doc
    wnDocFile = wnName + ".doc"
    wnDocObj = docx.Document()
    Log( "OutputName: " + wnFile + ', Doc: ' + wnDocFile )

    # Get webnovel url
    wnPage = input("Paste base URL here: ")
    Log( "Base URL: " + wnPage )

    # Get Title
    novelTitle = input( "Type Webnovel title: " )
    
    # Get webnovel details, from module ext
    novelInfo = modNovelInfo(wnPage)

    wnDocObj.add_paragraph("Title: " + novelTitle)
    wnDocObj.add_paragraph("Alt. Name: " + novelInfo[0])
    wnDocObj.add_paragraph("Author(s): " + novelInfo[1])
    wnDocObj.add_paragraph("Artist(s): " + novelInfo[2])
    wnDocObj.add_paragraph("Description: " + novelInfo[3])
    wnDocObj.add_paragraph("Year: " + novelInfo[4])
    wnDocObj.add_paragraph("Genre: " + novelInfo[5])

    # Get Chapter Links, from module ext
    listCh = modChapterLink(wnPage)

    # Counters
    count = 0
    countMax = len(listCh)
    status = ""

    # Iterate through chapter links
    for ch in listCh:
        # Clear screen
        clear()

        # Working on...
        count += 1
        status = "Working on " + str(count) + " out of " + str(countMax) + ".... Percentage: " + str((count/countMax)*100) + "%"
        print( status )
        Log( status )
        bodyString = ""
        titleString = ""

        # Get Body content, from module ext. And write to File
        body = modChInfo(ch)
        if len(body) < 1:
            bodyString = "NoneBody"
            titleString = "NoneTitle"
            Log( "Returns Empty Result" )
        else:
            # Get the title
            try:
                titleString = body[0]
                Log( "Has title" )
            except:
                Log( "No title contents" )
            # Get the chapter body content
            try:
                bodyString = body[1]
                Log( "Has contents" )
            except:
                Log( "No body contents" )

        # Content formatted
        contentToWrite = "Source: " + ch + "\nTitle: " + titleString +"\nBody:\n" + bodyString + "\n\n"

        # Write to Output File
        try:
            outputFile.write( contentToWrite )
            Log( "Written to output file!" )
        except:
            Log( "Not written to output! Encountered an error!" )
        
        # Add to docx file
        try:
            wnDocObj.add_paragraph( contentToWrite )
            Log( "Added to docx object!" )
        except:
            Log( "Not added to docx file! Encountered an error!" )

        # Done
        status = str(count) + " out of " + str(countMax) + " done! Percentage: " + str((count/countMax)*100) + "%"
        print( status )
        Log( status )

        time.sleep( 5.0 )
    # End of For Loop

    # Close output file
    outputFile.close()

    # Save to Doc file
    try:
        wnDocObj.save(wnDocFile)
        Log( "Written to docx file: " + wnDocFile )
    except:
        Log( "Cannot write to docx file! " + wnDocFile )

    print("Done!")

    Log( "##################################################################" )